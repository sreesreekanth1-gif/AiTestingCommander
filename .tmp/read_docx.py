from docx import Document
import sys

def read_docx(path):
    print("Reading", path)
    try:
        doc = Document(path)
        print("--- PARAGRAPHS ---")
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                print(f"[{i}] {p.text}")
        
        print("\n--- TABLES ---")
        for i, t in enumerate(doc.tables):
            print(f"Table {i}:")
            for j, row in enumerate(t.rows):
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(f"  Row {j}: {row_data}")
    except Exception as e:
        print("Error:", e)

if __name__ == "__main__":
    read_docx(sys.argv[1])
